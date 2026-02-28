#!/usr/bin/env python3
"""
Document generation utilities for HiMoM documentation project.
Converts markdown files to properly formatted .docx files with:
- Table of contents
- Styled headers, tables, and lists
- Page breaks between major sections
- Headers/footers with page numbers
- Screenshot placeholders styled distinctly
- Professional formatting for Google Docs compatibility
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def create_styled_document(title: str, subtitle: str = "") -> Document:
    """Create a new document with professional styling."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Style Heading 1
    h1 = doc.styles['Heading 1']
    h1.font.size = Pt(24)
    h1.font.color.rgb = RGBColor(0x1a, 0x56, 0x76)
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)

    # Style Heading 2
    h2 = doc.styles['Heading 2']
    h2.font.size = Pt(18)
    h2.font.color.rgb = RGBColor(0x2e, 0x74, 0x96)
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)

    # Style Heading 3
    h3 = doc.styles['Heading 3']
    h3.font.size = Pt(14)
    h3.font.color.rgb = RGBColor(0x3d, 0x85, 0xa8)
    h3.font.bold = True
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)

    # Title page
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(120)
    run = title_para.add_run(title)
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x1a, 0x56, 0x76)
    run.bold = True

    if subtitle:
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub_para.add_run(subtitle)
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Add date
    from datetime import date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.paragraph_format.space_before = Pt(24)
    run = date_para.add_run(f"Generated: {date.today().strftime('%B %d, %Y')}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Page break after title
    doc.add_page_break()

    # Add Table of Contents placeholder
    toc_heading = doc.add_paragraph('Table of Contents', style='Heading 1')

    # TOC field code
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = parse_xml(r'<w:fldChar {} w:fldCharType="begin"/>'.format(nsdecls('w')))
    run._r.append(fldChar)
    run = paragraph.add_run()
    instrText = parse_xml(r'<w:instrText {} xml:space="preserve"> TOC \o "1-3" \h \z \u</w:instrText>'.format(nsdecls('w')))
    run._r.append(instrText)
    run = paragraph.add_run()
    fldChar = parse_xml(r'<w:fldChar {} w:fldCharType="separate"/>'.format(nsdecls('w')))
    run._r.append(fldChar)
    run = paragraph.add_run('[Right-click and select "Update Field" to generate Table of Contents]')
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.italic = True
    run = paragraph.add_run()
    fldChar = parse_xml(r'<w:fldChar {} w:fldCharType="end"/>'.format(nsdecls('w')))
    run._r.append(fldChar)

    doc.add_page_break()

    # Add headers and footers
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = title
    header_para.style = doc.styles['Normal']
    header_run = header_para.runs[0] if header_para.runs else header_para.add_run(title)
    header_run.font.size = Pt(9)
    header_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Footer with page number
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run('Page ')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Page number field
    fldChar = parse_xml(r'<w:fldChar {} w:fldCharType="begin"/>'.format(nsdecls('w')))
    run._r.append(fldChar)
    run2 = footer_para.add_run()
    instrText = parse_xml(r'<w:instrText {} xml:space="preserve"> PAGE </w:instrText>'.format(nsdecls('w')))
    run2._r.append(instrText)
    run3 = footer_para.add_run()
    fldChar = parse_xml(r'<w:fldChar {} w:fldCharType="end"/>'.format(nsdecls('w')))
    run3._r.append(fldChar)

    return doc


def add_screenshot_placeholder(doc: Document, description: str):
    """Add a styled screenshot placeholder."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(12)

    # Add border styling via XML
    pPr = para._p.get_or_add_pPr()
    pBdr = parse_xml(
        '<w:pBdr {}>'
        '  <w:top w:val="single" w:sz="4" w:space="8" w:color="CCCCCC"/>'
        '  <w:left w:val="single" w:sz="4" w:space="8" w:color="CCCCCC"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="8" w:color="CCCCCC"/>'
        '  <w:right w:val="single" w:sz="4" w:space="8" w:color="CCCCCC"/>'
        '</w:pBdr>'.format(nsdecls('w'))
    )
    pPr.append(pBdr)

    # Shading
    shd = parse_xml('<w:shd {} w:fill="F5F5F5" w:val="clear"/>'.format(nsdecls('w')))
    pPr.append(shd)

    run = para.add_run(f"📷 SCREENSHOT: {description}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.font.italic = True


def add_info_box(doc: Document, text: str, box_type: str = "info"):
    """Add a styled info/warning/tip box."""
    colors = {
        "info": ("E8F4FD", "2196F3", "ℹ️"),
        "warning": ("FFF3E0", "FF9800", "⚠️"),
        "tip": ("E8F5E9", "4CAF50", "💡"),
        "error": ("FFEBEE", "F44336", "❌"),
    }
    bg, border, icon = colors.get(box_type, colors["info"])

    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(8)

    pPr = para._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg}" w:val="clear"/>')
    pPr.append(shd)

    run = para.add_run(f"{icon} {text}")
    run.font.size = Pt(10)


def add_styled_table(doc: Document, headers: list, rows: list):
    """Add a professionally styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    # Add spacing after table
    doc.add_paragraph()

    return table


def md_to_docx(md_path: str, docx_path: str, title: str, subtitle: str = ""):
    """Convert a markdown file to a styled .docx document."""
    md_content = Path(md_path).read_text(encoding='utf-8')

    doc = create_styled_document(title, subtitle)

    lines = md_content.split('\n')
    i = 0
    in_code_block = False
    in_table = False
    table_rows = []
    table_headers = []

    while i < len(lines):
        line = lines[i]

        # Code blocks
        if line.strip().startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue

        if in_code_block:
            para = doc.add_paragraph(line)
            para.style = doc.styles['Normal']
            for run in para.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            pPr = para._p.get_or_add_pPr()
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
            pPr.append(shd)
            i += 1
            continue

        # Tables
        if '|' in line and line.strip().startswith('|'):
            cells = [c.strip() for c in line.strip().strip('|').split('|')]

            if not in_table:
                # Check if next line is separator
                if i + 1 < len(lines) and re.match(r'\|[\s\-:|]+\|', lines[i + 1].strip()):
                    table_headers = cells
                    in_table = True
                    table_rows = []
                    i += 2  # Skip header and separator
                    continue
            else:
                if cells and any(c for c in cells):
                    table_rows.append(cells)
                else:
                    # End of table
                    if table_headers and table_rows:
                        add_styled_table(doc, table_headers, table_rows)
                    in_table = False
                    table_headers = []
                    table_rows = []
                i += 1
                continue
        elif in_table:
            # End of table (non-table line encountered)
            if table_headers and table_rows:
                add_styled_table(doc, table_headers, table_rows)
            in_table = False
            table_headers = []
            table_rows = []
            # Don't increment - process this line normally
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Screenshot placeholders
        screenshot_match = re.match(r'\[SCREENSHOT:\s*(.+?)\]', line.strip())
        if screenshot_match:
            add_screenshot_placeholder(doc, screenshot_match.group(1))
            i += 1
            continue

        # Headers
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue
        if line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue
        if line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=4)
            i += 1
            continue

        # Numbered lists
        num_match = re.match(r'^(\d+)\.\s+(.+)', line)
        if num_match:
            para = doc.add_paragraph(style='List Number')
            _add_formatted_text(para, num_match.group(2))
            i += 1
            continue

        # Bullet lists
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = re.sub(r'^[\s]*[-*]\s+', '', line)
            para = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(para, text)
            i += 1
            continue

        # Regular paragraph
        para = doc.add_paragraph()
        _add_formatted_text(para, line)
        i += 1

    # Flush any remaining table
    if in_table and table_headers and table_rows:
        add_styled_table(doc, table_headers, table_rows)

    doc.save(docx_path)
    print(f"Generated: {docx_path}")


def _add_formatted_text(para, text: str):
    """Add text with bold/italic formatting to a paragraph."""
    # Process bold and italic markers
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)

    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = para.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = para.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            # Light background for inline code
        else:
            para.add_run(part)


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: generate_docx.py <input.md> <output.docx> [title] [subtitle]")
        sys.exit(1)

    md_file = sys.argv[1]
    docx_file = sys.argv[2]
    title = sys.argv[3] if len(sys.argv) > 3 else Path(md_file).stem.replace('-', ' ').title()
    subtitle = sys.argv[4] if len(sys.argv) > 4 else "PreRollTracker & ApexAPI Documentation"

    md_to_docx(md_file, docx_file, title, subtitle)
